from django.http import HttpResponse

from rest_framework.permissions import IsAuthenticated
from rest_framework.decorators import api_view, permission_classes
from rest_framework.response import Response

import json
from io import BytesIO

from word.service.compare import (
    align_paragraphs_indexed,
    auto_thresholds,
    ratio,
    read_docx_lines_with_index,
)
from word.service.compare_rendering import (
    add_legend,
    write_sentence_aware_diff_with_ids,
)
from awcenter.file_security import WORD_POLICY, validate_request_upload

@api_view(["POST"])
@permission_classes([IsAuthenticated])
def compare(request):
    import pandas as pd
    from docx import Document

    file1 = validate_request_upload(request, "first", WORD_POLICY)
    file2 = validate_request_upload(request, "second", WORD_POLICY)
    parameters = json.loads(request.data["json"])

    equal_ratio = parameters["equal_ratio"]
    weak_equal_ratio = parameters["weak_equal_ratio"]
    output_type = parameters["output_type"]
    
    lines1 = read_docx_lines_with_index(file1)
    lines2 = read_docx_lines_with_index(file2)

    # Auto Ratio
    if equal_ratio is None or weak_equal_ratio is None:
        ar, wr = auto_thresholds(lines1, lines2)
        equal_ratio = ar if equal_ratio is None else equal_ratio
        weak_equal_ratio = wr if weak_equal_ratio is None else weak_equal_ratio

    aligned = align_paragraphs_indexed(lines1, lines2, equal_ratio, weak_equal_ratio)

    # DOCX
    if output_type == "word":
        doc = Document()
        doc.add_heading("Karşılaştırma Raporu (Gelişmiş Redline + Otomatik Eşik + ID)", level=1)
        p = doc.add_paragraph()
        p.add_run(f"Kaynak: {file1}  →  Hedef: {file2}").italic = True
        p = doc.add_paragraph()
        p.add_run(f"equal_ratio={equal_ratio:.3f}  weak_equal_ratio={weak_equal_ratio:.3f}").italic = True
        add_legend(doc)
        doc.add_paragraph()

    stats = {"equal": 0, "insert": 0, "delete": 0, "replace": 0}
    excel_rows = []

    for a_idx, a_text, b_idx, b_text, tag in aligned:
        stats[tag] += 1

        if output_type == "word": # DOCX content
            if tag == "equal":
                write_sentence_aware_diff_with_ids(doc, a_idx, a_text, b_idx, b_text) # To ignore equal lines, delete it
            else:
                write_sentence_aware_diff_with_ids(doc, a_idx, a_text, b_idx, b_text)
        elif output_type == "excel": # Excel content
            excel_rows.append({
                "Tag": tag,                    # equal / insert / delete / replace
                "A_Index": a_idx,
                "B_Index": b_idx,
                "A_Text": a_text,
                "B_Text": b_text,
                "A_Len": len(a_text or ""),
                "B_Len": len(b_text or ""),
                "Similarity": ratio(a_text or "", b_text or "") if (a_text and b_text) else None
            })

    if output_type == "word":
        # Summary Page
        doc.add_page_break()
        doc.add_heading("Özet / Summary", level=2)
        s = doc.add_paragraph()
        s.add_run(
            f"Eş (equal): {stats['equal']}  |  Eklenen (insert): {stats['insert']}  |  "
            f"Silinen (delete): {stats['delete']}  |  Değiştirilen (replace): {stats['replace']}"
        )

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        res = HttpResponse(buffer.getvalue(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        res["Content-Disposition"] = 'attachment; filename="Comparison Result.docx"'
        return res
        

    if output_type == "excel":
        # Excel
        df = pd.DataFrame(excel_rows, columns=[
            "Tag","A_Index","B_Index","A_Text","B_Text","A_Len","B_Len","Similarity"
        ])
        # Secondary summary table
        summary_df = pd.DataFrame([stats]).rename(columns={
            "equal":"Equal","insert":"Insert","delete":"Delete","replace":"Replace"
        })

        buffer = BytesIO()
        write_excel_report_openpyxl(
            out_excel=buffer,
            excel_rows=excel_rows,
            summary_stats=stats,
            table_rows=None
        )
        buffer.seek(0)

        res = HttpResponse(buffer.getvalue(), content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        res["Content-Disposition"] = 'attachment; filename="Comparison Result.xlsx"'
        return res

    return Response("Error while comparing files.", status=400)



def write_excel_report_openpyxl(out_excel, excel_rows, summary_stats, table_rows=None):
    import pandas as pd
    from openpyxl.formatting.rule import FormulaRule
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    # DataFrames
    diff_cols = [
        "Tag","A_Index","B_Index","A_Text","B_Text",
        "A_Len","B_Len","Similarity",
        "ListFlag","ListLevel_A","ListLevel_B","ListTag"
    ]
    
    # Automatically create if no fields exist
    if excel_rows and not any(("ListFlag" in r) for r in excel_rows):
        for r in excel_rows:
            r.setdefault("ListFlag", None)
            r.setdefault("ListLevel_A", None)
            r.setdefault("ListLevel_B", None)
            r.setdefault("ListTag", None)

    diff_df = pd.DataFrame(excel_rows, columns=diff_cols)

    summary_df = pd.DataFrame([{
        "Equal": summary_stats.get("equal", 0),
        "Insert": summary_stats.get("insert", 0),
        "Delete": summary_stats.get("delete", 0),
        "Replace": summary_stats.get("replace", 0),
        "Total": sum(summary_stats.values()) if summary_stats else len(diff_df)
    }])

    # TableDiff Page
    if table_rows:
        table_cols = ["TableIndex","ChangeType","RowKey","ColumnName","OldValue","NewValue"]
        table_df = pd.DataFrame(table_rows, columns=table_cols)
    else:
        table_df = None

    # Writer
    with pd.ExcelWriter(out_excel, engine="openpyxl") as writer:
        # Diff
        diff_df.to_excel(writer, index=False, sheet_name="Diff")
        ws = writer.sheets["Diff"]

        # Title freeze and filter
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions

        # Auto columns width
        for col_idx, col_name in enumerate(diff_df.columns, start=1):
            max_len = max(
                [len(str(col_name))] +
                [len(str(v)) if v is not None else 0 for v in diff_df[col_name].tolist()]
            )
            ws.column_dimensions[get_column_letter(col_idx)].width = min(60, max(10, max_len + 2))

        # Title style
        header_fill = PatternFill(start_color="FFDCE6F1", end_color="FFDCE6F1", fill_type="solid")
        for cell in ws[1]:
            cell.font = Font(bold=True)
            cell.fill = header_fill
            cell.alignment = Alignment(vertical="center")

        # Conditional style for tag column  (color)
        # Assuming A columns is Tag
        last_row = ws.max_row
        tag_range = f"A2:A{last_row}"

        # insert -> green text
        ws.conditional_formatting.add(
            tag_range,
            FormulaRule(formula=[f'ISNUMBER(SEARCH("insert",A2))'], stopIfTrue=False,
                        font=Font(color="FF008000"))
        )
        # delete -> red text
        ws.conditional_formatting.add(
            tag_range,
            FormulaRule(formula=[f'ISNUMBER(SEARCH("delete",A2))'], stopIfTrue=False,
                        font=Font(color="FFFF0000"))
        )
        # replace -> blue text
        ws.conditional_formatting.add(
            tag_range,
            FormulaRule(formula=[f'ISNUMBER(SEARCH("replace",A2))'], stopIfTrue=False,
                        font=Font(color="FF1F4E78"))
        )

        # Data Bar for Similarity
        sim_col_idx = None
        for c_idx, name in enumerate(diff_df.columns, start=1):
            if name == "Similarity":
                sim_col_idx = c_idx
                break
        if sim_col_idx:
            col_letter = get_column_letter(sim_col_idx)
            data_bar_range = f"{col_letter}2:{col_letter}{last_row}"
            ws.conditional_formatting.add(
                data_bar_range,
                DataBarRule(start_type="num", start_value=0,
                            end_type="num", end_value=1, color="FF63BE7B",
                            showValue=True)
            )

        # Summary
        summary_df.to_excel(writer, index=False, sheet_name="Summary")
        ws_sum = writer.sheets["Summary"]
        ws_sum.freeze_panes = "A2"
        ws_sum.auto_filter.ref = ws_sum.dimensions
        for col_idx, col_name in enumerate(summary_df.columns, start=1):
            max_len = max(
                [len(str(col_name))] +
                [len(str(v)) if v is not None else 0 for v in summary_df[col_name].tolist()]
            )
            ws_sum.column_dimensions[get_column_letter(col_idx)].width = min(40, max(10, max_len + 2))
        for cell in ws_sum[1]:
            cell.font = Font(bold=True)
            cell.fill = header_fill
            cell.alignment = Alignment(vertical="center")

        # TableDiff (optional)
        if table_df is not None and len(table_df) > 0:
            table_df.to_excel(writer, index=False, sheet_name="TableDiff")
            ws_tab = writer.sheets["TableDiff"]
            ws_tab.freeze_panes = "A2"
            ws_tab.auto_filter.ref = ws_tab.dimensions
            for col_idx, col_name in enumerate(table_df.columns, start=1):
                max_len = max(
                    [len(str(col_name))] +
                    [len(str(v)) if v is not None else 0 for v in table_df[col_name].tolist()]
                )
                ws_tab.column_dimensions[get_column_letter(col_idx)].width = min(60, max(10, max_len + 2))
            for cell in ws_tab[1]:
                cell.font = Font(bold=True)
                cell.fill = header_fill
                cell.alignment = Alignment(vertical="center")

            # add color to ChangeType column
            last_row_tab = ws_tab.max_row
            chg_col = None
            for i, n in enumerate(table_df.columns, start=1):
                if n == "ChangeType":
                    chg_col = get_column_letter(i)
                    break
            if chg_col:
                rng = f"{chg_col}2:{chg_col}{last_row_tab}"
                ws_tab.conditional_formatting.add(
                    rng,
                    FormulaRule(formula=[f'ISNUMBER(SEARCH("row-inserted",{chg_col}2))'],
                                font=Font(color="FF008000"))
                )
                ws_tab.conditional_formatting.add(
                    rng,
                    FormulaRule(formula=[f'ISNUMBER(SEARCH("row-deleted",{chg_col}2))'],
                                font=Font(color="FFFF0000"))
                )
                ws_tab.conditional_formatting.add(
                    rng,
                    FormulaRule(formula=[f'ISNUMBER(SEARCH("cell-changed",{chg_col}2))'],
                                font=Font(color="FF1F4E78"))
                )
