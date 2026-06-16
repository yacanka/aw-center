"""DOCX rendering helpers for Word comparison reports."""

from difflib import SequenceMatcher

from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt

from word.service.compare import ratio, split_sentences, tokenize_words

# DOCX redline style

def style_added(run):
    """Apply legacy added-text styling to a DOCX run."""
    run.font.underline = True
    run.font.color.rgb = None
    run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN  # optional

def style_deleted(run):
    """Apply legacy deleted-text styling to a DOCX run."""
    run.font.strike = True
    run.font.color.rgb = None
    run.font.highlight_color = WD_COLOR_INDEX.RED  # optional

def style_normal(run):
    """Apply normal text styling to a DOCX run."""
    run.font.strike = False
    run.font.underline = False

def add_legend(doc):
    """Add the comparison legend to a DOCX document."""
    p = doc.add_paragraph()
    r = p.add_run("Legend: ")
    r.bold = True
    r.font.size = Pt(11)
    p = doc.add_paragraph()
    r = p.add_run("Added")
    style_added(r)
    p.add_run("  |  ")
    r2 = p.add_run("Deleted")
    style_deleted(r2)
    p.add_run("  |  Changes are shown first as Deleted, then as Added.")

def write_word_diff(par, a_text: str, b_text: str):
    """Write a token-level redline diff into a DOCX paragraph."""
    a_tokens = tokenize_words(a_text or "")
    b_tokens = tokenize_words(b_text or "")
    sm = SequenceMatcher(None, a_tokens, b_tokens, autojunk=False)
    first = True
    def add(tok, styler):
        nonlocal first
        if not first:
            par.add_run(" ")
        r = par.add_run(tok)
        styler(r)
        first = False
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            for t in a_tokens[i1:i2]: add(t, style_normal)
        elif tag == "delete":
            for t in a_tokens[i1:i2]: add(t, style_deleted)
        elif tag == "insert":
            for t in b_tokens[j1:j2]: add(t, style_added)
        elif tag == "replace":
            for t in a_tokens[i1:i2]: add(t, style_deleted)
            for t in b_tokens[j1:j2]: add(t, style_added)

def write_sentence_aware_diff_with_ids(doc, a_idx, a_text, b_idx, b_text):
    """Write a sentence-aware diff block with source line identifiers."""
    hdr = doc.add_paragraph()
    hdr_run = hdr.add_run(f"[A#{a_idx if a_idx is not None else '-'} | B#{b_idx if b_idx is not None else '-'}]")
    hdr_run.italic = True

    # insert/delete
    if a_text is None and b_text is not None:
        par = doc.add_paragraph()
        for tok in tokenize_words(b_text):
            r = par.add_run(tok + " "); style_added(r)
        return

    if b_text is None and a_text is not None:
        par = doc.add_paragraph()
        for tok in tokenize_words(a_text):
            r = par.add_run(tok + " "); style_deleted(r)
        return

    # Aware word/sentence diff when both exist
    r = ratio(a_text, b_text)
    if r >= 0.95:
        par = doc.add_paragraph()
        write_word_diff(par, a_text, b_text)
        return

    s1 = split_sentences(a_text)
    s2 = split_sentences(b_text)
    sm = SequenceMatcher(None, s1, s2, autojunk=False)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            for k in range(i2 - i1):
                par = doc.add_paragraph()
                write_word_diff(par, s1[i1 + k], s2[j1 + k])
        elif tag == "delete":
            for k in range(i2 - i1):
                par = doc.add_paragraph()
                for tok in tokenize_words(s1[i1 + k]):
                    r = par.add_run(tok + " "); style_deleted(r)
        elif tag == "insert":
            for k in range(j2 - j1):
                par = doc.add_paragraph()
                for tok in tokenize_words(s2[j1 + k]):
                    r = par.add_run(tok + " "); style_added(r)
        elif tag == "replace":
            par = doc.add_paragraph()
            write_word_diff(par, " ".join(s1[i1:i2]), " ".join(s2[j1:j2]))

