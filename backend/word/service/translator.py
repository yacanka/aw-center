import re
from docx import Document
from io import BytesIO
import threading
from typing import Dict

_lock = threading.Lock()
_registry: Dict[str, "DocTranslator"] = {}

class DocTranslator:
    PLACEHOLDERS = {
        "\t": "XZY0",
        "\u00A0": "XZY1",  # non-breaking space
        "\u00AD": "",            # soft hyphen (çoğu zaman görünmez, sorun çıkarır)
    }

    REVERSE_PLACEHOLDERS = {v: k for k, v in PLACEHOLDERS.items() if v}
    
    MODELS = {
        "tr2en": r"models\opus-mt-tr-en",
        "en2tr": r"models\opus-mt-tc-big-en-tr"
    }
    
    def __init__(self, translate_type=None):
        from transformers import MarianMTModel, MarianTokenizer

        model_id = self.MODELS.get(translate_type)
        if not model_id:
            raise ValueError("Unsupported translate type.")
        
        self.tokenizer = MarianTokenizer.from_pretrained(model_id, local_files_only=True)
        self.model = MarianMTModel.from_pretrained(model_id, local_files_only=True)        

    def translate_text(self, text: str) -> str:
        text = text.strip()
        if not text:
            return text

        # Çok uzun metni parça parça çevirmek için basit bölme:
        # (512 token sınırı var; karakter bazlı bölme pratikte iş görüyor)
        max_chars = 900  # güvenli
        if len(text) <= max_chars:
            return self._translate_one(text)

        parts = self.split_preserving_separators(text, max_chars=max_chars)
        out = []
        for p in parts:
            if p.strip():
                out.append(self._translate_one(p))
            else:
                out.append(p)
        return "".join(out)

    def _translate_one(self, text: str) -> str:
        tokens = self.tokenizer(text, return_tensors="pt", truncation=True, max_length=512)
        translated = self.model.generate(**tokens)
        return self.tokenizer.decode(translated[0], skip_special_tokens=True)

    def split_preserving_separators(self, text: str, max_chars: int = 900):
        """
        Metni olabildiğince cümle sonlarından/paragraf içi ayrımlardan bölmeye çalışır.
        Ayraçları da (boşluk/satır sonu) korur.
        """
        if len(text) <= max_chars:
            return [text]

        # Önce satır sonlarından böl
        chunks = []
        current = ""
        for piece in re.split(r"(\n+)", text):
            if len(current) + len(piece) <= max_chars:
                current += piece
            else:
                if current:
                    chunks.append(current)
                current = piece
        if current:
            chunks.append(current)

        # Hâlâ büyük parça varsa cümle sonlarından böl
        final = []
        for ch in chunks:
            if len(ch) <= max_chars:
                final.append(ch)
                continue
            cur = ""
            for piece in re.split(r"([.!?…]+[\s\n]*)", ch):
                if len(cur) + len(piece) <= max_chars:
                    cur += piece
                else:
                    if cur:
                        final.append(cur)
                    cur = piece
            if cur:
                final.append(cur)

        # Yine büyük kalırsa sert böl
        really_final = []
        for ch in final:
            if len(ch) <= max_chars:
                really_final.append(ch)
            else:
                for i in range(0, len(ch), max_chars):
                    really_final.append(ch[i:i+max_chars])
        return really_final


    # -----------------------------------
    # 2) RUN MERGE -> TRANSLATE -> REDISTRIBUTE
    # -----------------------------------

    # Çeviri sırasında bozulmaması için bazı karakterleri token’a çevirip geri koyuyoruz.
    

    def protect_specials(self, s: str) -> str:
        for k, v in self.PLACEHOLDERS.items():
            s = s.replace(k, v)
        return s

    def restore_specials(self, s: str) -> str:
        # soft hyphen'i geri koymuyoruz (bilerek)
        for k, v in self.REVERSE_PLACEHOLDERS.items():
            s = s.replace(k, v)
        return s

    def looks_like_field_or_toc(self, text: str) -> bool:
        """
        Word alanları (TOC, PAGE, DATE vb.) python-docx’te tam görünmese de
        bazı şablonlarda bu metinler run.text olarak gelebiliyor.
        Riskli görünenleri çevirme.
        """
        t = text.strip().upper()
        if not t:
            return True
        # TOC/field benzeri tipik anahtarlar
        field_markers = ["TOC", "PAGEREF", "HYPERLINK", "PAGE", "NUMPAGES", "DATE", "REF "]
        return any(m in t for m in field_markers)

    def redistribute_by_original_lengths(self, translated: str, run_lengths: list[int]) -> list[str]:
        """
        Çevrilen metni, orijinal run uzunluk oranlarına göre bölüp geri dağıtır.
        - Kelime parçalamayı azaltmak için önce hedef aralıkta en yakın boşluğu arar.
        - Eğer hiç boşluk yoksa sert keser.
        """
        total = sum(run_lengths)
        if total <= 0:
            return [translated] + [""] * (len(run_lengths) - 1)

        # Hedef kesim noktaları (kümülatif)
        cut_targets = []
        acc = 0
        for L in run_lengths[:-1]:
            acc += L
            cut_targets.append(acc / total)

        # Çevrilen metin üzerinde kesimler
        N = len(translated)
        cuts = []
        last = 0

        for ratio in cut_targets:
            target_idx = int(round(ratio * N))

            # hedefin çevresinde boşluk arayalım
            best = None
            window = 25  # ne kadar yakın arayalım
            lo = max(last, target_idx - window)
            hi = min(N, target_idx + window)

            # sağa doğru boşluk bulma öncelikli (kelimeyi bölmemek için)
            for i in range(target_idx, hi):
                if translated[i:i+1].isspace():
                    best = i + 1
                    break
            if best is None:
                # sola doğru bak
                for i in range(target_idx, lo, -1):
                    if translated[i-1:i].isspace():
                        best = i
                        break

            if best is None:
                best = target_idx

            cuts.append(best)

        parts = []
        prev = 0
        for c in cuts:
            parts.append(translated[prev:c])
            prev = c
        parts.append(translated[prev:])

        # Parça sayısı run sayısı ile aynı olmalı
        if len(parts) < len(run_lengths):
            parts += [""] * (len(run_lengths) - len(parts))
        elif len(parts) > len(run_lengths):
            # Fazlalığı sona birleştir
            extra = "".join(parts[len(run_lengths)-1:])
            parts = parts[:len(run_lengths)-1] + [extra]

        return parts

    def translate_paragraph_preserve_runs(self, paragraph):
        runs = paragraph.runs
        if not runs:
            return

        # Çevrilecek run’ları seç (sadece metin içerenler)
        texts = [r.text for r in runs]
        if not any(t.strip() for t in texts):
            return

        # Field/TOC şüphesi varsa pas geç (çok agresif istemiyorsan bu satırı kapatabilirsin)
        joined_raw = "".join(texts)
        if self.looks_like_field_or_toc(joined_raw):
            return

        # Koruma: tab, nbsp vb.
        protected = self.protect_specials(joined_raw)

        # Çevir
        translated = self.translate_text(protected)
        translated = self.restore_specials(translated)

        # Run uzunlukları (sadece text uzunluğu; boş run’lar da dahil olsun ki stil yapısı korunabilsin)
        run_lengths = [len(t) for t in texts]
        if sum(run_lengths) == 0:
            # hepsi boşsa
            return

        parts = self.redistribute_by_original_lengths(translated, run_lengths)

        # Yaz
        for r, new_text in zip(runs, parts):
            r.text = new_text

    def iter_all_paragraphs(self, doc: Document):
        # Ana gövde paragrafları
        total_paragraphs = len(doc.paragraphs)
        for i, p in enumerate(doc.paragraphs):
            yield (p, i, total_paragraphs, "paragraphs")

        # Tabloların içi
        total_tables = len(doc.tables)
        for i, table in enumerate(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield (p, i, total_tables, "tables")

        # Header / Footer (section bazlı)
        
        for section in doc.sections:
            total_section_header_p = len(section.header.paragraphs)
            for i, p in enumerate(section.header.paragraphs):
                yield (p, i, total_section_header_p, "header paragraphs")
            
            total_section_header_t = len(section.header.tables)
            for i, table in enumerate(section.header.tables):
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield (p, i, total_section_header_t, "header tables")

            total_section_footer_p = len(section.footer.paragraphs)
            for i, p in enumerate(section.footer.paragraphs):
                yield (p, i, total_section_footer_p, "footer paragraphs")
                
            total_section_footer_t = len(section.footer.tables)
            for i, table in enumerate(section.footer.tables):
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield (p, i, total_section_footer_t, "footer tables")

    def translate_docx_file(self, input_path: str, output_path: str):
        doc = Document(input_path)
        for p, i, total, ptype in self.iter_all_paragraphs(doc):
            self.translate_paragraph_preserve_runs(p)
        doc.save(output_path)
    
    def translate_docx_req(self, input_bytes: BytesIO):
        doc = Document(input_bytes)
        for p, i, total, ptype in self.iter_all_paragraphs(doc):
            self.translate_paragraph_preserve_runs(p)
            yield ("progress", (i, total, ptype))

        buffer = BytesIO()
        doc.save(buffer)
        yield ("result", buffer)


def get_text_generator(generator_type) -> DocTranslator:
    dt = _registry.get(generator_type)
    if dt is None:
        with _lock:
            dt = _registry.get(generator_type)
            if dt is None:
                dt = DocTranslator(generator_type)
                _registry[generator_type] = dt
    return dt
