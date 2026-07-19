from io import BytesIO
from unittest.mock import patch
from zipfile import ZipFile

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import override_settings
from docx import Document
from openpyxl import Workbook

from excel.cover_pages import REQUIRED_COLUMNS
from jobs.models import Job, JobStatus
from jobs.services import create_job
from jobs.worker import claim_next_job, execute_claimed_job
from .base import JobTestCase


class DocumentJobApiTests(JobTestCase):
    """Verify durable Word and Excel enqueue contracts."""

    def test_word_translation_job_accepts_only_safe_direction(self):
        """Translation direction is allowlisted before job persistence."""

        response = self.client.post(
            "/word/jobs/translate/",
            {"file": word_upload(), "translate_type": "tr2en"},
            format="multipart",
        )

        self.assertEqual(response.status_code, 201)
        job = Job.objects.get(pk=response.data["id"])
        self.assertEqual(job.kind, "word.translate")
        self.assertEqual(job.parameters, {"translate_type": "tr2en"})

    def test_cover_page_job_rejects_missing_columns(self):
        """Invalid workbook contracts fail before consuming worker capacity."""

        response = self.client.post(
            "/excel/jobs/cover-pages/",
            {"file": excel_upload(["Wrong Column"])},
            format="multipart",
        )

        self.assertEqual(response.status_code, 400)
        self.assertEqual(Job.objects.count(), 0)

    def test_cover_page_job_never_persists_unrelated_session_parameters(self):
        """The cover-page contract stores no JIRA session or unused browser state."""

        response = self.client.post(
            "/excel/jobs/cover-pages/",
            {"file": excel_upload(REQUIRED_COLUMNS), "JSESSIONID": "secret-session"},
            format="multipart",
        )

        self.assertEqual(response.status_code, 201)
        self.assertEqual(Job.objects.get(pk=response.data["id"]).parameters, {})


class DocumentJobWorkerTests(JobTestCase):
    """Verify Word and Excel executors produce durable artifacts."""

    @patch("word.job_executor.get_text_generator")
    def test_word_executor_persists_translated_document(self, generator_mock):
        """Translator progress and result become a completed job artifact."""

        generator_mock.return_value.translate_docx_req.return_value = translated_results()
        job, _ = create_job(
            self.user, "word.translate", "Translate", {"translate_type": "tr2en"}, word_upload()
        )

        execute_claimed_job(claim_next_job("word-worker"))

        job.refresh_from_db()
        self.assertEqual(job.status, JobStatus.SUCCEEDED)
        self.assertTrue(job.output_name.endswith(".docx"))
        self.assertGreater(job.output_file.size, 0)

    @patch("word.job_executor.get_text_generator", side_effect=OSError("private path"))
    def test_word_executor_sanitizes_missing_model_failure(self, _generator_mock):
        """Model paths and raw loader failures never reach the job contract."""

        job, _ = create_job(
            self.user, "word.translate", "Translate", {"translate_type": "tr2en"}, word_upload()
        )

        execute_claimed_job(claim_next_job("word-worker"))

        job.refresh_from_db()
        self.assertEqual(job.status, JobStatus.FAILED)
        self.assertEqual(job.error_code, "WORD_MODEL_UNAVAILABLE")
        self.assertTrue(job.retryable)
        self.assertNotIn("private path", job.message)
        detail = self.client.get(f"/jobs/{job.id}/")
        self.assertIn("deploy", detail.data["recovery_hint"].lower())

    @override_settings(COVER_PAGE_TEMPLATE_PATH="/missing/template.docx")
    def test_excel_executor_uses_safe_builtin_fallback(self):
        """Cover pages remain functional when no custom template is deployed."""

        job, _ = create_job(
            self.user, "excel.cover_pages", "Cover pages", {}, excel_upload(REQUIRED_COLUMNS)
        )

        execute_claimed_job(claim_next_job("excel-worker"))

        job.refresh_from_db()
        self.assertEqual(job.status, JobStatus.SUCCEEDED)
        with job.output_file.open("rb") as output, ZipFile(output) as archive:
            names = archive.namelist()
            self.assertEqual(len(names), 1)
            self.assertNotIn("..", names[0])
            self.assertTrue(names[0].endswith(".docx"))


def word_upload():
    """Return a decoder-valid Word upload."""

    buffer = BytesIO()
    document = Document()
    document.add_paragraph("Merhaba dünya")
    document.save(buffer)
    return SimpleUploadedFile(
        "document.docx", buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def excel_upload(columns):
    """Return a decoder-valid workbook with one representative row."""

    buffer = BytesIO()
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.append(columns)
    worksheet.append(["B30/001" if column == "Cover Page Number" else "value" for column in columns])
    workbook.save(buffer)
    return SimpleUploadedFile(
        "cover-pages.xlsx", buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


def translated_results():
    """Return a deterministic translator generator result."""

    buffer = BytesIO()
    document = Document()
    document.add_paragraph("Hello world")
    document.save(buffer)
    return iter([("progress", (1, 1, "paragraph")), ("result", buffer)])
