import json
import zipfile
from types import SimpleNamespace
from unittest.mock import Mock, patch

from django.core.files.base import ContentFile
from docx import Document
from awcenter.job_executors import resolve_job_executor
from dcc.document_job import build_render_context
from dcc.document_snapshot import DccSnapshotError, build_snapshot, capture_dcc_snapshot
from dcc.document_preview import prepare_dcc_preview
from dcc.services.template_resolver import DccTemplateNotFoundError
from jobs.models import Job, JobStatus
from jobs.services import create_job
from jobs.worker import claim_next_job, execute_claimed_job
from orgs.models import Project, ProjectRoleAssignment

from .base import JobTestCase


class DccSnapshotTests(JobTestCase):
    """Verify DCC snapshot boundary cases independently of JIRA transport."""

    @patch("dcc.job_views.jira_connector_for")
    def test_preview_and_confirmation_reject_query_string_credentials(self, connector_for):
        """JIRA credentials cannot be smuggled through durable DCC job URLs."""

        preview = self.client.post(
            "/api/dcc/jobs/create-document/preview/?JSESSIONID=never-in-a-url",
            {"url": "DCC-1"},
            format="json",
        )
        confirmation = self.client.post(
            "/api/dcc/jobs/create-document/00000000-0000-0000-0000-000000000001/confirm/"
            "?JSESSIONID=never-in-a-url",
            {},
            format="json",
        )

        self.assertEqual(preview.status_code, 400)
        self.assertEqual(confirmation.status_code, 400)
        self.assertEqual(preview.data["code"], "JIRA_SESSION_CANONICAL_REQUIRED")
        self.assertEqual(confirmation.data["code"], "JIRA_SESSION_CANONICAL_REQUIRED")
        connector_for.assert_not_called()

    def test_zero_subtasks_produces_a_valid_snapshot(self):
        """Parent tasks without panels no longer fail on an undefined last subtask."""

        issue = parent_issue([])
        project = Project.objects.get(slug="hys")
        snapshot = build_snapshot(
            fake_connector(), issue, "DCC-1", [project_definition()], [project]
        )

        self.assertEqual(snapshot["issue_key"], "DCC-1")
        self.assertEqual(snapshot["schema_version"], 2)
        self.assertEqual(snapshot["placeholders"]["Panels"], [])
        self.assertNotIn("Responsible_AS", snapshot["placeholders"])

    def test_panel_placeholders_are_an_ordered_array_without_numbered_keys(self):
        connector = fake_connector([panel_issue()])
        issue = parent_issue([SimpleNamespace(key="P-1")])
        project = Project.objects.get(slug="hys")

        snapshot = build_snapshot(
            connector, issue, "DCC-1", [project_definition()], [project]
        )

        placeholders = snapshot["placeholders"]
        self.assertEqual(
            placeholders["Panels"][0]["Panel_AS_Name"],
            "Ada LOVELACE",
        )
        self.assertEqual(placeholders["Panels"][0]["Panel_Status"], "Done")
        self.assertFalse(any(key.endswith("_1") for key in placeholders))

    def test_project_controller_changes_only_the_pre_render_context(self):
        connector = fake_connector([panel_issue()])
        issue = parent_issue([SimpleNamespace(key="P-1")])
        project = Project.objects.get(slug="gokbey")
        snapshot = build_snapshot(
            connector,
            issue,
            "DCC-1",
            [project_definition("gokbey")],
            [project],
        )

        context = build_render_context(snapshot)

        self.assertEqual(
            snapshot["placeholders"]["Panels"][0]["Panel_AS_Name"],
            "Ada LOVELACE",
        )
        self.assertEqual(
            context["Panels"][0]["Panel_AS_Name"],
            "Utku İnanç PEHLİVAN, Ada LOVELACE",
        )

    def test_conflicting_explicit_responsible_values_are_rejected(self):
        """Conflicting compliance owners fail visibly instead of selecting one silently."""

        panels = [panel_issue("Owner A"), panel_issue("Owner B")]
        connector = fake_connector(panels)
        issue = parent_issue([SimpleNamespace(key="P-1"), SimpleNamespace(key="P-2")])
        project = Project.objects.get(slug="hys")

        with self.assertRaises(DccSnapshotError) as raised:
            build_snapshot(
                connector, issue, "DCC-1", [project_definition()], [project]
            )

        self.assertEqual(raised.exception.code, "DCC_RESPONSIBLE_CONFLICT")

    def test_capture_rejects_multiple_resolved_projects_before_panel_reads(self):
        hys = Project.objects.get(slug="hys")
        ProjectRoleAssignment.objects.create(
            user=self.user,
            project=hys,
            domain=ProjectRoleAssignment.Domain.DCC,
            role=ProjectRoleAssignment.Role.OPERATOR,
        )
        issue = parent_issue([])
        issue.fields.components = [
            SimpleNamespace(name="HYS"),
            SimpleNamespace(name="GOKBEY"),
        ]
        connector = Mock()
        connector.current_user.return_value = {"name": "jira-user"}
        connector.get_issue.return_value = issue
        connector.get_client.return_value.issue.side_effect = AssertionError(
            "Panel reads must not happen before authorization."
        )

        with self.assertRaises(DccSnapshotError) as raised:
            capture_dcc_snapshot(connector, "DCC-1", self.user)

        self.assertEqual(raised.exception.code, "DCC_PROJECT_AMBIGUOUS")
        connector.get_client.assert_not_called()


class DccDocumentExecutorTests(JobTestCase):
    """Exercise real DOCX rendering, validation, and private artifact persistence."""

    def test_confirmed_preview_reaches_verified_completion(self):
        """The reviewed API flow must continue from confirmation through the worker."""

        template_path = self.media_directory / "dcc-template.docx"
        create_template(template_path)
        preview, confirmation = self.complete_confirmed_preview(template_path)

        self.assertEqual(preview.status_code, 201)
        self.assertEqual(confirmation.data["status"], JobStatus.QUEUED)
        persisted_status = Job.objects.values_list("status", flat=True).get(pk=preview.data["id"])
        self.assertEqual(persisted_status, JobStatus.SUCCEEDED)

    def complete_confirmed_preview(self, template_path):
        """Run a real preview, confirmation, and worker execution."""

        project = Project.objects.get(slug="hys")
        ProjectRoleAssignment.objects.create(
            user=self.user,
            project=project,
            domain=ProjectRoleAssignment.Domain.DCC,
            role=ProjectRoleAssignment.Role.OPERATOR,
        )
        with (
            patch(
                "dcc.job_views.capture_dcc_snapshot",
                return_value=snapshot_contract([project.pk]),
            ),
            patch("dcc.job_views.jira_connector_for"),
            patch("dcc.document_job.get_project_definition", return_value=project_definition()),
            patch("dcc.document_job.resolve_dcc_template_path", return_value=template_path),
        ):
            preview = self.client.post(
                "/api/dcc/jobs/create-document/preview/",
                {"url": "DCC-1"}, format="json",
                HTTP_IDEMPOTENCY_KEY="dcc-complete-flow",
            )
            warning_codes = preview.data["result_summary"]["readiness_warning_codes"]
            confirmation = self.client.post(
                f"/api/dcc/jobs/create-document/{preview.data['id']}/confirm/",
                {"acknowledged_warning_codes": warning_codes}, format="json",
            )
            execute_claimed_job(
                claim_next_job("dcc-complete-worker"), resolve_job_executor
            )
        return preview, confirmation

    def test_worker_renders_a_real_verified_docx(self):
        """The allowlisted worker produces a readable OOXML document without base64."""

        template_path = self.media_directory / "dcc-template.docx"
        create_template(template_path)
        upload = ContentFile(json.dumps(snapshot_contract()).encode(), name="dcc-DCC-1.json")
        job, _created = create_job(
            self.user, "dcc.create_document", "Create DCC", {"issue_key": "DCC-1"}, upload
        )

        with patch("dcc.document_job.get_project_definition", return_value=project_definition()):
            with patch("dcc.document_job.resolve_dcc_template_path", return_value=template_path):
                execute_claimed_job(claim_next_job("dcc-worker"), resolve_job_executor)

        job.refresh_from_db()
        self.assertEqual(job.status, JobStatus.SUCCEEDED)
        with job.output_file.open("rb") as output:
            output_path = self.media_directory / "result.docx"
            output_path.write_bytes(output.read())
        self.assertTrue(zipfile.is_zipfile(output_path))
        output_text = "\n".join(p.text for p in Document(output_path).paragraphs)
        self.assertIn("Change title", output_text)
        self.assertIn("Ada LOVELACE", output_text)

    def test_schema_v1_numbered_panel_fields_are_upgraded_for_queued_jobs(self):
        snapshot = snapshot_contract()
        snapshot["schema_version"] = 1
        snapshot["panel_count"] = 1
        snapshot["placeholders"] = {
            "Panel_AS_Name_1": "Ada LOVELACE",
            "Panel_Status_1": "Done",
        }

        context = build_render_context(snapshot)

        self.assertEqual(
            context["Panels"],
            [{"Panel_AS_Name": "Ada LOVELACE", "Panel_Status": "Done"}],
        )
        self.assertNotIn("Panel_AS_Name_1", context)

    def test_preview_dry_renders_exact_snapshot_and_reports_omissions(self):
        """Preflight proves template readiness and exposes only safe missing-field labels."""

        template_path = self.media_directory / "dcc-template.docx"
        create_template(template_path)
        snapshot = snapshot_contract()

        with patch("dcc.document_job.get_project_definition", return_value=project_definition()):
            with patch("dcc.document_job.resolve_dcc_template_path", return_value=template_path):
                summary = prepare_dcc_preview(snapshot)

        self.assertTrue(summary["template_ready"])
        self.assertEqual(summary["output_name"], "DCC-1.docx")
        self.assertIn("DCC form number", summary["missing_recommended_fields"])
        self.assertEqual(summary["readiness_level"], "review")
        self.assertTrue(summary["requires_readiness_acknowledgement"])
        self.assertNotIn("Change title", str(summary))

    def test_missing_template_has_a_retryable_stable_failure(self):
        """Deployment template failures remain recoverable without exposing paths."""

        upload = ContentFile(json.dumps(snapshot_contract()).encode(), name="dcc-DCC-1.json")
        job, _created = create_job(
            self.user, "dcc.create_document", "Create DCC", {"issue_key": "DCC-1"}, upload
        )
        with patch("dcc.document_job.resolve_dcc_template_path") as resolver:
            resolver.side_effect = DccTemplateNotFoundError("private/path/template.docx")
            execute_claimed_job(claim_next_job("dcc-worker"), resolve_job_executor)

        job.refresh_from_db()
        self.assertEqual(job.status, JobStatus.FAILED)
        self.assertEqual(job.error_code, "DCC_TEMPLATE_UNAVAILABLE")
        self.assertTrue(job.retryable)
        self.assertNotIn("private/path", job.message)


def snapshot_contract(project_ids=None):
    """Return a minimal versioned rendering snapshot."""

    return {
        "schema_version": 2,
        "issue_key": "DCC-1",
        "project_slug": "hys",
        "project_slugs": ["hys"],
        "project_ids": list(project_ids or []),
        "project_label": "HYS",
        "output_name": "DCC-1.docx",
        "panel_count": 1,
        "placeholders": {
            "Design_Change_Title": "Change title",
            "Panels": [{"Panel_AS_Name": "Ada LOVELACE"}],
        },
    }


def parent_issue(subtasks):
    """Return a minimal parent issue accepted by the snapshot builder."""

    fields = SimpleNamespace(
        subtasks=subtasks, summary="Change title", customfield_45002="DCC-1",
        customfield_45000="ECD-1", customfield_45001="A", customfield_13716=None,
        updated="2026-06-30T00:00:00+00:00", customfield_34115=[],
    )
    return SimpleNamespace(fields=fields)


def panel_issue(responsible=""):
    """Return a minimal panel subtask issue."""

    fields = SimpleNamespace(
        summary="Flight Panel Assessment", status=SimpleNamespace(name="Done"),
        assignee=SimpleNamespace(displayName="Ada Lovelace"),
        updated="2026-06-30T00:00:00+00:00", customfield_45006=None,
        customfield_45007=None, customfield_45008="Assessment", customfield_45421=None,
        customfield_45004=SimpleNamespace(value="Minor-No Effect"),
        customfield_45005=responsible, comment=SimpleNamespace(comments=[]),
    )
    return SimpleNamespace(fields=fields)


def fake_connector(panels=None):
    """Return a connector double serving panel issues in order."""

    iterator = iter(panels or [])
    return SimpleNamespace(get_client=lambda: SimpleNamespace(issue=lambda _key: next(iterator)))


def project_definition(slug="hys"):
    """Return the DCC-capable project fields required by snapshot and renderer."""

    label = slug.upper()
    return SimpleNamespace(slug=slug, dcc_label=label, jira_component=label)


def create_template(path):
    """Create a real DOCX template with one DocxTemplate placeholder."""

    document = Document()
    document.add_paragraph("{{ Design_Change_Title }}")
    document.add_paragraph(
        "{% for panel in Panels %}{{ panel.Panel_AS_Name }};{% endfor %}"
    )
    document.save(path)
