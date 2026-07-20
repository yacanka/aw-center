"""Regression tests for the CompDoc-to-DCC traceability bridge."""

import json
from uuid import uuid4
from unittest.mock import patch

from django.contrib.auth.models import Permission
from docx import Document

from dcc.compdoc_bridge import (
    MAX_COMPDOC_SELECTION,
    attach_compliance_documents,
    parse_compdoc_selection,
)
from dcc.compdoc_register import append_compdoc_register
from dcc.document_snapshot import DccSnapshotError
from jobs.models import Job
from jobs.tests.base import JobTestCase
from projects.ozgur.models import CompDoc

PREVIEW_URL = "/dcc/jobs/create-document/preview/"


class CompdocBridgeTests(JobTestCase):
    """Verify authorization, immutable capture, API persistence, and DOCX output."""

    def setUp(self):
        super().setUp()
        self.document = CompDoc.objects.create(
            name="Flight Manual", cover_page_no="CP-001", cover_page_issue="A",
            tech_doc_no="FM-100", tech_doc_issue="2", panel="Flight", responsible="Owner",
            status_flow=[{"status": "authority_review", "date": "20.07.2026"}],
        )
        self.dcc_permission = permission("dcc", "add_jira_dcc")
        self.view_permission = permission("ozgur", "view_compdoc")
        self.user.user_permissions.add(self.dcc_permission)

    def test_bridge_requires_project_compdoc_view_permission(self):
        """DCC creation permission alone cannot disclose project compliance data."""

        with self.assertRaises(DccSnapshotError) as raised:
            attach_compliance_documents(snapshot(), self.user, "ozgur", [str(self.document.pk)])

        self.assertEqual(raised.exception.code, "DCC_COMPDOC_FORBIDDEN")
        self.assertEqual(raised.exception.response_status, 403)

    def test_capture_is_versioned_deterministic_and_immutable(self):
        """Each capture binds exact history versions and changes digest after an update."""

        self.user.user_permissions.add(self.view_permission)
        first = attach_compliance_documents(
            snapshot(), self.user, "ozgur", [str(self.document.pk)]
        )["compliance_documents"]
        self.document.name = "Updated Flight Manual"
        self.document.save()
        second = attach_compliance_documents(
            snapshot(), self.user, "ozgur", [str(self.document.pk)]
        )["compliance_documents"]

        self.assertEqual(first["documents"][0]["name"], "Flight Manual")
        self.assertNotEqual(first["fingerprint"], second["fingerprint"])
        self.assertNotEqual(
            first["documents"][0]["source_history_id"],
            second["documents"][0]["source_history_id"],
        )

    @patch("dcc.job_views.prepare_dcc_preview")
    @patch("dcc.job_views.capture_dcc_snapshot")
    def test_preview_persists_compdoc_sources_without_credentials(self, capture, prepare):
        """The owner-bound job stores references and content, never the JIRA session."""

        self.user.user_permissions.add(self.view_permission)
        capture.return_value = snapshot()
        prepare.return_value = preview_summary()
        response = self.client.post(
            PREVIEW_URL,
            preview_payload(self.document.pk, session="secret-session"),
            format="json",
            HTTP_IDEMPOTENCY_KEY="compdoc-dcc-preview",
        )

        self.assertEqual(response.status_code, 201)
        job = Job.objects.get(pk=response.data["id"])
        self.assertEqual(job.parameters["compdoc_ids"], [str(self.document.pk)])
        with job.input_file.open("rb") as source:
            payload = json.load(source)
        self.assertEqual(payload["compliance_documents"]["documents"][0]["name"], "Flight Manual")
        self.assertNotIn("secret-session", json.dumps(payload))

    @patch("dcc.job_views.capture_dcc_snapshot", return_value=None)
    def test_invalid_selection_fails_before_jira_capture(self, capture):
        """Malformed and excessive selections never cause an external JIRA read."""

        payload = preview_payload(self.document.pk)
        payload["compdoc_ids"] = [str(uuid4())] * (MAX_COMPDOC_SELECTION + 1)
        response = self.client.post(PREVIEW_URL, payload, format="json")

        self.assertEqual(response.status_code, 400)
        self.assertEqual(response.data["code"], "DCC_COMPDOC_LIMIT")
        capture.assert_not_called()

    @patch("dcc.job_views.prepare_dcc_preview")
    @patch("dcc.job_views.capture_dcc_snapshot")
    def test_confirmation_rechecks_revoked_compdoc_access(self, capture, prepare):
        """Permission revocation blocks queued use of an existing private snapshot."""

        self.user.user_permissions.add(self.view_permission)
        capture.return_value = snapshot()
        prepare.return_value = preview_summary()
        preview = self.client.post(
            PREVIEW_URL, preview_payload(self.document.pk), format="json",
            HTTP_IDEMPOTENCY_KEY="compdoc-revocation-preview",
        )
        self.user.user_permissions.remove(self.view_permission)
        self.user = type(self.user).objects.get(pk=self.user.pk)
        self.client.force_authenticate(self.user)

        response = self.client.post(
            f"/dcc/jobs/create-document/{preview.data['id']}/confirm/", format="json"
        )

        self.assertEqual(response.status_code, 403)
        self.assertEqual(response.data["code"], "DCC_COMPDOC_FORBIDDEN")
        self.assertEqual(Job.objects.get(pk=preview.data["id"]).status, "awaiting_confirmation")

    def test_register_appends_human_readable_provenance(self):
        """Generated DCC files contain selected references and their SHA-256 fingerprint."""

        self.user.user_permissions.add(self.view_permission)
        bundle = attach_compliance_documents(
            snapshot(), self.user, "ozgur", [str(self.document.pk)]
        )["compliance_documents"]
        output_path = self.media_directory / "register.docx"
        Document().save(output_path)

        append_compdoc_register(output_path, bundle)

        rendered = Document(output_path)
        text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
        cells = [cell.text for row in rendered.tables[0].rows for cell in row.cells]
        self.assertIn(bundle["fingerprint"], text)
        self.assertIn("Flight Manual", cells)
        self.assertIn("FM-100 / 2", cells)


class CompdocSelectionValidationTests(JobTestCase):
    """Protect the bounded request contract independently of database records."""

    def test_selection_rejects_duplicates_and_partial_context(self):
        """A selection must be unique and paired with exactly one project."""

        identifier = str(uuid4())
        with self.assertRaises(DccSnapshotError) as duplicate:
            parse_compdoc_selection({"compdoc_project": "ozgur", "compdoc_ids": [identifier] * 2})
        with self.assertRaises(DccSnapshotError) as partial:
            parse_compdoc_selection({"compdoc_ids": [identifier]})

        self.assertEqual(duplicate.exception.code, "DCC_COMPDOC_DUPLICATE")
        self.assertEqual(partial.exception.code, "DCC_COMPDOC_SELECTION_INVALID")


def permission(app_label, codename):
    """Return an unambiguous model permission fixture."""

    return Permission.objects.get(content_type__app_label=app_label, codename=codename)


def snapshot():
    """Return a minimal Ozgur DCC rendering snapshot."""

    return {
        "schema_version": 1, "issue_key": "DCC-1", "project_slug": "ozgur",
        "project_label": "Ozgur", "output_name": "DCC-1.docx", "panel_count": 0,
        "placeholders": {"Design_Change_Title": "Change"},
    }


def preview_payload(document_id, session="temporary"):
    """Return a complete CompDoc-backed DCC preview request."""

    return {
        "JSESSIONID": session, "url": "DCC-1", "compdoc_project": "ozgur",
        "compdoc_ids": [str(document_id)],
    }


def preview_summary():
    """Return safe mocked preview metadata."""

    return {
        "type": "dcc_preview", "issue_key": "DCC-1", "project": "Ozgur",
        "output_name": "DCC-1.docx", "panel_count": 0, "template_ready": True,
    }
