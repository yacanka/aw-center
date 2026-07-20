"""Append an immutable compliance-document register to generated DCC files."""

from docx import Document

REGISTER_HEADERS = ("Document", "Cover Page", "Technical Document", "Panel / ATA", "Status", "Responsible")


def append_compdoc_register(output_path, bundle):
    """Append selected CompDocs to a rendered DOCX without template changes."""

    documents = bundle.get("documents", []) if isinstance(bundle, dict) else []
    if not documents:
        return
    document = Document(output_path)
    document.add_page_break()
    document.add_heading("Compliance Document Traceability Register", level=1)
    add_provenance(document, bundle)
    table = document.add_table(rows=1, cols=len(REGISTER_HEADERS))
    apply_table_style(table)
    fill_row(table.rows[0].cells, REGISTER_HEADERS)
    for source in documents:
        fill_row(table.add_row().cells, register_values(source))
    document.save(output_path)


def apply_table_style(table):
    """Prefer a bordered table while tolerating minimal project style catalogs."""

    try:
        table.style = "Table Grid"
    except KeyError:
        pass


def add_provenance(document, bundle):
    """Add content-free capture provenance above the register."""

    fingerprint = str(bundle.get("fingerprint") or "")
    project = str(bundle.get("project_slug") or "")
    captured_at = str(bundle.get("captured_at") or "")
    document.add_paragraph(f"Project: {project} | Captured: {captured_at}")
    document.add_paragraph(f"Source fingerprint (SHA-256): {fingerprint}")


def register_values(source):
    """Return one human-readable traceability table row."""

    cover_page = join_reference(source.get("cover_page_no"), source.get("cover_page_issue"))
    technical = "; ".join(
        join_reference(item.get("number"), item.get("issue"))
        for item in source.get("technical_documents", [])
    )
    panel = " / ".join(filter(None, (source.get("panel"), source.get("ata"))))
    return (source.get("name"), cover_page, technical, panel, source.get("status"), source.get("responsible"))


def join_reference(number, issue):
    """Combine a document number and issue without manufacturing values."""

    parts = [str(value).strip() for value in (number, issue) if value]
    return " / ".join(parts)


def fill_row(cells, values):
    """Fill a DOCX table row using plain text only."""

    for cell, value in zip(cells, values):
        cell.text = str(value or "")
