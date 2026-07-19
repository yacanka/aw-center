import re
from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from django.conf import settings
from docx import Document
from docxtpl import DocxTemplate

from jobs.artifacts import materialize_job_input, temporary_output
from jobs.contracts import JobExecutionFailure, JobExecutionResult
from jobs.worker import update_progress

REQUIRED_COLUMNS = [
    "Cover Page Number", "Cover Page Issue", "ATA Chapter", "Disciplines",
    "Technical Document Name", "CAT", "Requirements", "MoC",
    "Technical Document Number", "Technical Document Issue", "AS Name", "CVE Name",
]


def inspect_workbook_columns(file_source):
    """Return canonical column names and missing required fields."""

    import pandas as pd
    try:
        frame = pd.read_excel(file_source, nrows=0)
        columns = [str(column).strip() for column in frame.columns]
        missing = [column for column in REQUIRED_COLUMNS if column.casefold() not in casefolded(columns)]
        return columns, missing
    except Exception as error:
        raise ValueError("The Excel workbook could not be read.") from error
    finally:
        if hasattr(file_source, "seek"):
            file_source.seek(0)


def execute_cover_page_creation(job):
    """Create a ZIP of cover-page documents from one validated workbook."""

    input_path = materialize_job_input(job)
    output_path = temporary_output(".zip")
    result_ready = False
    try:
        rows = load_cover_rows(input_path)
        create_cover_page_archive(job.id, rows, output_path)
        result_ready = True
        return JobExecutionResult(output_path, "cover-pages.zip", "Cover pages created.")
    finally:
        input_path.unlink(missing_ok=True)
        if not result_ready:
            output_path.unlink(missing_ok=True)


def load_cover_rows(input_path):
    """Load, canonicalize, and bound cover-page workbook rows."""

    import pandas as pd
    try:
        frame = pd.read_excel(input_path)
    except Exception as error:
        raise JobExecutionFailure("The Excel workbook could not be read.", "EXCEL_INVALID") from error
    frame = canonicalize_columns(frame)
    frame = frame.dropna(subset=["Cover Page Number", "Cover Page Issue"], how="any").fillna("")
    if frame.empty:
        raise JobExecutionFailure("No complete cover-page rows were found.", "COVER_PAGE_ROWS_EMPTY")
    if len(frame) > settings.COVER_PAGE_MAX_ROWS:
        raise JobExecutionFailure("The workbook exceeds the cover-page row limit.", "COVER_PAGE_ROW_LIMIT")
    return frame


def canonicalize_columns(frame):
    """Map case-insensitive workbook columns to the canonical contract."""

    lookup = {str(column).strip().casefold(): column for column in frame.columns}
    missing = [column for column in REQUIRED_COLUMNS if column.casefold() not in lookup]
    if missing:
        raise JobExecutionFailure(
            f"Missing required columns: {', '.join(missing)}", "COVER_PAGE_COLUMNS_MISSING"
        )
    return frame.rename(columns={lookup[column.casefold()]: column for column in REQUIRED_COLUMNS})


def create_cover_page_archive(job_id, rows, output_path):
    """Render each workbook row into a bounded ZIP archive."""

    template_bytes = load_template_bytes()
    total = len(rows)
    project_label = resolve_project_label(str(rows.iloc[0]["Cover Page Number"]))
    with ZipFile(output_path, mode="w", compression=ZIP_DEFLATED) as archive:
        for index, (_, source_row) in enumerate(rows.iterrows(), start=1):
            row = build_template_row(source_row, project_label)
            archive.writestr(cover_filename(row, index), render_cover_page(row, template_bytes))
            progress = 10 + int(index / total * 80)
            update_progress(job_id, progress, f"Created cover page {index}/{total}.")
            enforce_archive_limit(output_path)


def build_template_row(source_row, project_label):
    """Convert canonical columns into template-safe lowercase placeholders."""

    row = {str(key).strip().lower().replace(" ", "_"): str(value) for key, value in source_row.items()}
    row["requirements"] = row["requirements"].strip().replace("\n", ", ").replace(",, ", ", ")
    row["project_name"] = project_label
    return row


def render_cover_page(row, template_bytes):
    """Render a trusted custom template or a safe built-in fallback."""

    buffer = BytesIO()
    if template_bytes:
        template = DocxTemplate(BytesIO(template_bytes))
        template.render(row, autoescape=True)
        template.save(buffer)
    else:
        built_in_cover_page(row).save(buffer)
    return buffer.getvalue()


def built_in_cover_page(row):
    """Create a readable fallback cover page without an external template."""

    document = Document()
    document.add_heading("Compliance Document Cover Page", level=1)
    document.add_paragraph(row["project_name"])
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in row.items():
        cells = table.add_row().cells
        cells[0].text = label.replace("_", " ").title()
        cells[1].text = str(value)
    return document


def load_template_bytes():
    """Load the configured trusted template when it exists."""

    template_path = Path(settings.COVER_PAGE_TEMPLATE_PATH)
    return template_path.read_bytes() if template_path.is_file() else None


def cover_filename(row, index):
    """Build a traversal-safe and duplicate-safe archive entry name."""

    identifier = re.sub(r"[^A-Za-z0-9._-]+", "_", row["cover_page_number"]).strip("._")
    return f"{index:04d}-{identifier[:80] or 'cover-page'}.docx"


def resolve_project_label(cover_page_number):
    """Resolve a configured project label from a cover-page identifier."""

    if "B30" in cover_page_number.upper():
        return settings.COVER_PAGE_B30_PROJECT_LABEL
    return settings.COVER_PAGE_DEFAULT_PROJECT_LABEL


def enforce_archive_limit(output_path):
    """Stop archive growth before it consumes unbounded private storage."""

    if output_path.stat().st_size > settings.JOB_MAX_OUTPUT_BYTES:
        raise JobExecutionFailure("Generated archive exceeds the safety limit.", "JOB_OUTPUT_TOO_LARGE")


def casefolded(values):
    """Return a case-insensitive set of string values."""

    return {value.casefold() for value in values}
